"""Document Ingest Pipeline - PDFからテキスト・テーブル・画像を抽出"""

import fitz  # PyMuPDF
import pdfplumber
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from loguru import logger
import re
from PIL import Image
import io


class DocumentIngestor:
    """入札書類からデータを抽出"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.xlsx']

    def ingest(self, file_path: str) -> Dict[str, Any]:
        """
        メインの取り込み処理

        Args:
            file_path: ドキュメントのパス

        Returns:
            抽出されたデータ
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported format: {path.suffix}")

        logger.info(f"Ingesting document: {path.name}")

        if path.suffix.lower() == '.pdf':
            return self._ingest_pdf(str(path))
        elif path.suffix.lower() == '.docx':
            return self._ingest_docx(str(path))
        elif path.suffix.lower() == '.xlsx':
            return self._ingest_excel(str(path))

        raise ValueError(f"Handler not implemented for: {path.suffix}")

    def _ingest_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        PDFからテキスト、テーブル、画像を抽出

        Returns:
            {
                'text': 全テキスト,
                'pages': ページごとのデータ,
                'tables': 抽出されたテーブル,
                'images': 画像データ,
                'metadata': PDFメタデータ
            }
        """
        result = {
            'text': '',
            'pages': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }

        # PyMuPDFでテキストと画像を抽出
        doc = fitz.open(file_path)
        result['metadata'] = {
            'page_count': len(doc),
            'title': doc.metadata.get('title', ''),
            'author': doc.metadata.get('author', ''),
            'subject': doc.metadata.get('subject', ''),
        }

        all_text = []

        for page_num, page in enumerate(doc, start=1):
            page_data = {
                'page_number': page_num,
                'text': '',
                'tables': [],
                'has_images': False
            }

            # テキスト抽出
            text = page.get_text()
            page_data['text'] = text
            all_text.append(text)

            # 画像抽出
            image_list = page.get_images()
            if image_list:
                page_data['has_images'] = True
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    try:
                        base_image = doc.extract_image(xref)
                        result['images'].append({
                            'page': page_num,
                            'index': img_index,
                            'ext': base_image['ext'],
                            'size': len(base_image['image'])
                        })
                    except Exception as e:
                        logger.warning(f"Failed to extract image on page {page_num}: {e}")

            result['pages'].append(page_data)

        doc.close()

        result['text'] = '\n\n'.join(all_text)

        # pdfplumberでテーブルを抽出
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    for table_index, table in enumerate(tables):
                        if table:
                            result['tables'].append({
                                'page': page_num,
                                'index': table_index,
                                'data': table,
                                'rows': len(table),
                                'cols': len(table[0]) if table else 0
                            })

                            # ページデータにもテーブル情報を追加
                            if page_num <= len(result['pages']):
                                result['pages'][page_num - 1]['tables'].append({
                                    'index': table_index,
                                    'data': table
                                })
        except Exception as e:
            logger.error(f"Failed to extract tables: {e}")

        logger.info(f"Extracted {len(result['pages'])} pages, {len(result['tables'])} tables, {len(result['images'])} images")

        return result

    def _ingest_docx(self, file_path: str) -> Dict[str, Any]:
        """Word文書からデータを抽出"""
        from docx import Document

        doc = Document(file_path)

        result = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {}
        }

        # パラグラフ抽出
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                result['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name
                })

        result['text'] = '\n\n'.join(paragraphs)

        # テーブル抽出
        for table_index, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            result['tables'].append({
                'index': table_index,
                'data': table_data,
                'rows': len(table_data),
                'cols': len(table_data[0]) if table_data else 0
            })

        logger.info(f"Extracted {len(paragraphs)} paragraphs, {len(result['tables'])} tables from DOCX")

        return result

    def _ingest_excel(self, file_path: str) -> Dict[str, Any]:
        """Excelファイルからデータを抽出"""
        import pandas as pd

        result = {
            'text': '',
            'sheets': [],
            'tables': [],
            'metadata': {}
        }

        # 全シートを読み込み
        excel_file = pd.ExcelFile(file_path)
        result['metadata']['sheet_names'] = excel_file.sheet_names

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # DataFrameをリストに変換
            table_data = [df.columns.tolist()] + df.values.tolist()

            result['sheets'].append({
                'name': sheet_name,
                'data': table_data,
                'rows': len(df),
                'cols': len(df.columns)
            })

            result['tables'].append({
                'sheet': sheet_name,
                'data': table_data,
                'rows': len(df),
                'cols': len(df.columns)
            })

        logger.info(f"Extracted {len(result['sheets'])} sheets from Excel")

        return result

    def extract_project_info(self, ingested_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        抽出されたデータから案件情報を抽出

        Returns:
            {
                'project_name': 案件名,
                'client_name': 顧客名,
                'location': 所在地,
                etc.
            }
        """
        text = ingested_data.get('text', '')

        project_info = {}

        # 案件名を抽出（様々なパターンに対応）
        patterns = {
            'project_name': [
                r'件\s*名[：:]\s*(.+)',
                r'工事名[：:]\s*(.+)',
                r'【(.+)】',
            ],
            'client_name': [
                r'発注者[：:]\s*(.+)',
                r'契約者[：:]\s*(.+)',
            ],
            'location': [
                r'所在地[：:]\s*(.+)',
                r'場\s*所[：:]\s*(.+)',
            ],
            'contract_period': [
                r'契約期間[：:]\s*(.+)',
                r'賃貸借期間[：:]\s*(.+)',
            ],
        }

        for key, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, text)
                if match:
                    project_info[key] = match.group(1).strip()
                    break

        return project_info

    def extract_building_specs(self, ingested_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        建物仕様を抽出（諸元表から）

        Returns:
            建物仕様のリスト
        """
        text = ingested_data.get('text', '')
        tables = ingested_data.get('tables', [])

        building_specs = []

        # テーブルから諸元表を探す
        for table_info in tables:
            table_data = table_info['data']

            # 諸元表かどうかを判定
            if self._is_spec_table(table_data):
                specs = self._parse_spec_table(table_data)
                building_specs.extend(specs)

        # テキストから建物情報を抽出
        # 延床面積
        area_match = re.search(r'延[べ]*床面積[：:]\s*([\d,]+\.?\d*)\s*㎡', text)
        if area_match and building_specs:
            building_specs[0]['total_area'] = float(area_match.group(1).replace(',', ''))

        # 構造
        structure_match = re.search(r'構\s*造[：:]\s*(.+)', text)
        if structure_match and building_specs:
            building_specs[0]['structure'] = structure_match.group(1).strip()

        return building_specs

    def _is_spec_table(self, table_data: List[List[str]]) -> bool:
        """テーブルが諸元表かどうかを判定"""
        if not table_data or len(table_data) < 2:
            return False

        # ヘッダー行をチェック
        header = table_data[0] if table_data else []
        header_text = ' '.join(str(cell) for cell in header if cell).lower()

        # 諸元表のキーワードをチェック
        keywords = ['室名', '部屋', '面積', '設備', '仕様']
        return any(kw in header_text for kw in keywords)

    def _parse_spec_table(self, table_data: List[List[str]]) -> List[Dict[str, Any]]:
        """諸元表をパース"""
        if not table_data:
            return []

        specs = []
        header = table_data[0]

        # ヘッダーからカラムインデックスを特定
        col_mapping = {}
        for i, cell in enumerate(header):
            cell_text = str(cell).strip() if cell else ''
            if '室名' in cell_text or '部屋' in cell_text:
                col_mapping['room_name'] = i
            elif '面積' in cell_text:
                col_mapping['area'] = i
            elif '設備' in cell_text or '備考' in cell_text:
                col_mapping['equipment'] = i

        # データ行を処理
        for row in table_data[1:]:
            if not row or not any(row):
                continue

            spec = {}

            if 'room_name' in col_mapping and len(row) > col_mapping['room_name']:
                spec['room_name'] = str(row[col_mapping['room_name']]).strip()

            if 'area' in col_mapping and len(row) > col_mapping['area']:
                area_str = str(row[col_mapping['area']]).replace(',', '').strip()
                try:
                    spec['area'] = float(area_str)
                except (ValueError, AttributeError):
                    pass

            if 'equipment' in col_mapping and len(row) > col_mapping['equipment']:
                equipment_str = str(row[col_mapping['equipment']]).strip()
                spec['equipment'] = [e.strip() for e in equipment_str.split('、') if e.strip()]

            if spec.get('room_name'):
                specs.append(spec)

        return specs
